from docx import Document
from docx.shared import Inches
import os

def add_images_from_directory(doc, directory):
    # Obtém uma lista ordenada dos arquivos de imagem
    image_files = sorted([filename for filename in os.listdir(directory) if filename.lower().endswith(('.jpg', '.png'))])

    for filename in image_files:
        image_path = os.path.join(directory, filename)
        doc.add_paragraph()  # Adiciona um parágrafo em branco para ancorar a imagem
        run = doc.paragraphs[-1].add_run()
        run.add_picture(image_path, width=Inches(5))  # Adicione sua imagem aqui e ajuste o tamanho conforme necessário

# Carrega o documento existente
diretorio_template = r'resource\template\template.docx'
doc = Document(diretorio_template)

# Diretório onde estão as imagens
diretorio_imagens = r'logs\screenshot'

# Adicione as imagens do diretório ao documento
add_images_from_directory(doc, diretorio_imagens)

# Salve o documento modificado
doc.save('Evidencia_de_Teste.docx')
